from docx import Document
from docx.shared import Pt, RGBColor

class InformePsicologico():
    def __init__(self, template_path, file_path):
        self.doc = Document(template_path)
        self.file_path = file_path

    def replace(self, data: dict):
        """
        Reemplazar los valores en el documento Word manteniendo el estilo
        """
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            variable = "var_" + key
                            # variable = key
                            if variable in paragraph.text:
                                for run in paragraph.runs:
                                    font = run.font
                                    if variable in run.text:
                                        run.text = run.text.replace(variable, str(value))
                                        font = run.font
                                        if 'bold' in paragraph.style.name:
                                            font.bold = True
                                        else:
                                            font.bold = False
                                        font.size = Pt(12)  # Establecer el tamaño de la fuente según tus necesidades
                                        font.name = 'Arial'  # Establecer el tipo de fuente según tus necesidades

        for paragraph in self.doc.paragraphs:
            for key, value in data.items():
                variable = "var_" + key
                if variable in paragraph.text:
                    for run in paragraph.runs:
                        if variable in run.text:
                            run.text = run.text.replace(variable, str(value))
                            font = run.font
                            if 'bold' in paragraph.style.name:
                                font.bold = True
                            else:
                                font.bold = False
                            font.size = Pt(12)  # Establecer el tamaño de la fuente según tus necesidades
                            font.name = 'Arial'  # Establecer el tipo de fuente según tus necesidades

    def save(self):
        # Guardar el documento modificado con un nuevo nombre
        self.doc.save(self.file_path)
